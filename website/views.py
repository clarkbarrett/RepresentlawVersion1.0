from django.shortcuts import render, get_object_or_404
from django.views.generic import TemplateView
from blog.models import Post
# from docx import *
from docx import Document
# from docx.shared import Inches
from django.http import HttpResponse
import datetime
from io import StringIO


class HomeView(TemplateView):
    template_name = 'website/home.html'


def home(request, template_name):
    return render(request, template_name)


def latest_post_footer(request):
    post = Post.objects.order_by('-published')[0:3]
    template = 'website/base.html'
    context = {'posts': post}
    return render(request, template, context)

# Test
from reportlab.pdfgen import canvas
from django.http import HttpResponse


def some_view(request):
    if request.GET.get('pdf'):
        # Create the HttpResponse object with the appropriate PDF headers.
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="somefilename.pdf"'

        # Create the PDF object, using the response object as its "file."
        p = canvas.Canvas(response)

        # Draw things on the PDF. Here's where the PDF generation happens.
        # See the ReportLab documentation for the full list of functionality.
        p.drawString(100, 100, "Hello world.")

        # Close the PDF object cleanly, and we're done.
        p.showPage()
        p.save()
        return response

# End Test

'''
def test_document(request):

    document = Document()
    docx_title = "TEST_DOCUMENT.docx"
    # ---- Cover Letter ----
    # document.add_picture((r'%s/static/images/my-header.png' % (settings.PROJECT_PATH)), width=Inches(4))
    document.add_paragraph()
    document.add_paragraph("%s" % datetime.date.today().strftime('%d, %B, %Y'))

    document.add_paragraph('Dear Sir or Madam:')
    document.add_paragraph('We are pleased to help you with your widgets.')
    document.add_paragraph('Please feel free to contact me for any additional information.')
    document.add_paragraph('I look forward to assisting you in this project.')

    document.add_paragraph()
    document.add_paragraph('Best regards,')
    document.add_paragraph('Acme Specialist 1]')
    document.add_page_break()

    # Prepare document for download
    # -----------------------------
    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response'''
